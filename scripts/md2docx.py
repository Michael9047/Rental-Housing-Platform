"""将维修系统架构 MD 文档转为 Word"""
from docx import Document
from docx.shared import Pt
import os

doc = Document()

md_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'REPAIR_SYSTEM_ARCHITECTURE.md')

with open(md_path, encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    i += 1

    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(line.strip('*'))
        run.bold = True
    elif line.startswith('```'):
        # skip code blocks
        while i < len(lines) and not lines[i].startswith('```'):
            i += 1
        i += 1
    elif line.startswith('| '):
        # Read table
        rows = []
        while i < len(lines) and lines[i].startswith('| '):
            cells = [c.strip() for c in lines[i].split('|')[1:-1]]
            rows.append(cells)
            i += 1
        # Skip separator row (|---|)
        if len(rows) > 1 and all('-' in c for c in rows[0]):
            rows = rows[1:]
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]), style='Light Grid Accent 1')
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    text = cell.strip()
                    # Bold header
                    if ri == 0:
                        p = table.rows[ri].cells[ci].paragraphs[0]
                        run = p.add_run(text)
                        run.bold = True
                    else:
                        table.rows[ri].cells[ci].text = text
    elif line.startswith('---'):
        pass
    elif line.startswith('> '):
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        run.italic = True
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('  - '):
        doc.add_paragraph(line[4:], style='List Bullet')
    elif line.strip():
        doc.add_paragraph(line.strip())

# Save
docx_path = os.path.join(os.path.dirname(__file__), '..', 'docs', 'REPAIR_SYSTEM_ARCHITECTURE.docx')
doc.save(docx_path)
print(f'Done: {docx_path}')
